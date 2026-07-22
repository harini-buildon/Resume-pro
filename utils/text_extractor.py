"""
utils/text_extractor.py – Resume Text Extraction
==================================================
This module extracts text content from uploaded resume files.
It supports two formats:
  1. PDF  → uses the 'pdfplumber' library
  2. DOCX → uses the 'python-docx' library

KEY CONCEPTS FOR BEGINNERS:
──────────────────────────
- PDF files are NOT plain text. They store text in a complex binary format
  with positioning info, fonts, images, etc. We need a library to decode them.
  
- DOCX files are actually ZIP archives containing XML files.
  python-docx reads the XML and extracts the paragraph text for us.
  
- We use a "dispatcher" pattern: extract_text() checks the file type
  and calls the appropriate extraction function. This keeps the code clean.

ERROR HANDLING:
──────────────
- Corrupted PDFs → caught by try/except, returns an error message
- Empty files → detected and reported
- Password-protected PDFs → will fail gracefully
"""

import pdfplumber
from docx import Document


def extract_from_pdf(filepath):
    """
    Extract text from a PDF file using pdfplumber.
    
    How pdfplumber works:
    1. Opens the PDF file
    2. Iterates through each page
    3. Calls page.extract_text() to get the text content
    4. Joins all pages with newlines
    
    Parameters:
        filepath (str): Path to the PDF file
    
    Returns:
        tuple: (success: bool, result: str)
            - On success: (True, extracted_text)
            - On failure: (False, error_message)
    """
    try:
        text_parts = []
        
        # 'with' statement ensures the file is properly closed after reading
        with pdfplumber.open(filepath) as pdf:
            # Iterate through every page in the PDF
            for page_num, page in enumerate(pdf.pages, 1):
                page_text = page.extract_text()
                
                if page_text:
                    text_parts.append(page_text)
        
        # Join all pages into a single string
        full_text = '\n'.join(text_parts)
        
        # Check if we actually extracted any text
        if not full_text.strip():
            return False, "The PDF appears to be empty or contains only images (no extractable text)."
        
        return True, full_text
    
    except Exception as e:
        return False, f"Failed to extract text from PDF: {str(e)}"


def extract_from_docx(filepath):
    """
    Extract text from a DOCX (Word) file using python-docx.
    
    How python-docx works:
    1. Opens the DOCX file (which is a ZIP of XML files internally)
    2. Reads all paragraph elements
    3. Each paragraph has a .text attribute with the plain text
    
    Parameters:
        filepath (str): Path to the DOCX file
    
    Returns:
        tuple: (success: bool, result: str)
    """
    try:
        doc = Document(filepath)
        
        # Extract text from all paragraphs
        # doc.paragraphs is a list of Paragraph objects
        text_parts = []
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():  # Skip empty paragraphs
                text_parts.append(paragraph.text)
        
        # Also extract text from tables (resumes sometimes use tables for layout)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        text_parts.append(cell.text)
        
        full_text = '\n'.join(text_parts)
        
        if not full_text.strip():
            return False, "The DOCX file appears to be empty."
        
        return True, full_text
    
    except Exception as e:
        return False, f"Failed to extract text from DOCX: {str(e)}"


def extract_text(filepath):
    """
    Main dispatcher function – detects file type and extracts text.
    
    This is the function you'll call from app.py. It figures out
    whether the file is a PDF or DOCX and calls the right extractor.
    
    Parameters:
        filepath (str): Path to the uploaded resume file
    
    Returns:
        tuple: (success: bool, result: str)
    """
    # Determine the file type from the extension
    file_extension = filepath.rsplit('.', 1)[1].lower() if '.' in filepath else ''
    
    if file_extension == 'pdf':
        return extract_from_pdf(filepath)
    elif file_extension == 'docx':
        return extract_from_docx(filepath)
    else:
        return False, f"Unsupported file format: .{file_extension}. Please upload a PDF or DOCX file."
